import logging
import os
import sys
import re
import zipfile
import pandas as pd
from typing import List, Dict, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from aiohttp import web
import asyncio


from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.utils.keyboard import ReplyKeyboardMarkup, KeyboardButton
from aiogram.types import Update


from aiogram import Bot, Dispatcher, F, Router, types
from aiogram.filters import Command
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
from aiogram.utils.keyboard import ReplyKeyboardMarkup, KeyboardButton
from aiogram.types import FSInputFile, BotCommand, ErrorEvent, Update
from aiogram.fsm.storage.memory import MemoryStorage

# --- Логирование ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# --- Настройки ---
BOT_TOKEN = os.getenv("BOT_TOKEN")
WEBHOOK_URL = os.getenv("WEBHOOK_URL")  # полный URL, например https://yourdomain.com/webhook
WEBHOOK_PATH = "/webhook"
PORT = int(os.getenv("PORT", 10000))

if not BOT_TOKEN or not WEBHOOK_URL:
    logger.error("BOT_TOKEN или WEBHOOK_URL не установлен!")
    sys.exit(1)

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(storage=MemoryStorage())
router = Router()
dp.include_router(router)

# --- Состояния ---
class GenDocs(StatesGroup):
    choosing_stage = State()
    waiting_excel = State()
    waiting_template = State()

# --- Клавиатура для стадий ---
stage_keyboard = ReplyKeyboardMarkup(
    keyboard=[[KeyboardButton(text="ПД"), KeyboardButton(text="РД")]],
    resize_keyboard=True
)

# --- Обработчики ---
@router.message(Command("start"))
async def start_cmd(message: types.Message, state: FSMContext):
    await message.answer("Привет! Для какой стадии нужно сделать титульные листы?", reply_markup=stage_keyboard)
    await state.set_state(GenDocs.choosing_stage)

@router.message(GenDocs.choosing_stage, F.text.in_(["ПД", "РД"]))
async def choose_stage(message: types.Message, state: FSMContext):
    stage = message.text.strip()
    await state.update_data(stage=stage)
    await message.answer(f"Вы выбрали стадию: {stage}. Теперь отправьте Excel файл 📑")
    await state.set_state(GenDocs.waiting_excel)

@router.message(GenDocs.waiting_excel, F.document)
async def handle_excel(message: types.Message, state: FSMContext):
    data = await state.get_data()
    stage = data.get("stage")

    file = await bot.get_file(message.document.file_id)

    # --- Проверка расширения ---
    if not file.file_path.endswith(".xlsx"):
        await message.answer("⚠️ Пожалуйста, отправьте Excel-файл в формате .xlsx")
        return

    file_path = f"{stage}_data.xlsx"
    await bot.download_file(file.file_path, file_path)

    # --- Проверка чтения Excel ---
    try:
        if stage == "ПД":
            _ = pd.read_excel(file_path, usecols=[0, 1, 2])
        else:
            _ = pd.read_excel(file_path)
    except Exception as e:
        logger.error(f"Ошибка чтения Excel: {e}")
        await message.answer("⚠️ Не удалось прочитать Excel. Проверьте, что файл корректный.")
        if os.path.exists(file_path):
            os.remove(file_path)
        return

    await state.update_data(excel_path=file_path)
    await message.answer("Файл Excel получен ✅ Теперь отправьте шаблон Word (.docx)")
    await state.set_state(GenDocs.waiting_template)


@router.message(GenDocs.waiting_template, F.document)
async def handle_template(message: types.Message, state: FSMContext):
    data = await state.get_data()
    stage = data.get("stage")
    excel_path = data.get("excel_path")

    file = await bot.get_file(message.document.file_id)

    if not file.file_path.endswith(".docx"):
        await message.answer("⚠️ Пожалуйста, отправьте Word-шаблон в формате .docx")
        return

    template_path = f"{stage}_template.docx"
    await bot.download_file(file.file_path, template_path)

    try:
        _ = Document(template_path)
    except Exception as e:
        logger.error(f"Ошибка открытия Word шаблона: {e}")
        await message.answer("⚠️ Не удалось открыть шаблон Word. Проверьте, что файл корректный.")
        if os.path.exists(template_path):
            os.remove(template_path)
        return

    try:
        if stage == "ПД":
            df = pd.read_excel(excel_path, usecols=[0, 1, 2])
            subtables = split_dataframe_PD(df)
            archive_name = create_word_for_each_row_PD(subtables, template_path, f"{stage}_docs.zip")
        else:
            df = pd.read_excel(excel_path)
            subtable = split_dataframe_RD(df)
            archive_name = create_word_for_each_row_RD(subtable, template_path, f"{stage}_docs.zip")

        await message.answer("Документы сгенерированы ✅ Вот ваш архив:")
        await message.answer_document(FSInputFile(archive_name))

    except Exception as e:
        logger.error(f"Ошибка генерации документов: {e}")
        await message.answer("⚠️ Произошла ошибка при генерации документов.")
    finally:
        # Очистка временных файлов
        for file_path in [excel_path, template_path, f"{stage}_docs.zip"]:
            if os.path.exists(file_path):
                os.remove(file_path)
        
        await state.clear()

@router.message(Command("help"))
async def help_cmd(message: types.Message):
    help_text = """
🤖 **Инструкция по использованию бота:**

1. Нажмите /start
2. Выберите стадию (ПД или РД)
3. Отправьте Excel файл с данными
4. Отправьте Word шаблон

📊 **Требования к Excel файлу:**
- Для ПД: должны быть колонки Том, Шифр, Часть
- Для РД: должны быть колонки Шифр, Раздел

📝 **Word шаблон должен содержать поля для замены:**
- Для ПД: Номер, Название шифра, Название части, Название раздела, Название подраздела
- Для РД: Название шифра, Название раздела
"""
    await message.answer(help_text)

    try:
        await message.answer_document(FSInputFile("examples/Состав_ПД.xlsx"))
        await message.answer_document(FSInputFile("examples/Титул_ПД.docx"))
        await message.answer_document(FSInputFile("examples/Состав_РД.xlsx"))
        await message.answer_document(FSInputFile("examples/Титул_РД.docx"))
    except Exception as e:
        logger.error(f"Ошибка отправки примеров: {e}")
        await message.answer("⚠️ Примеры недоступны. Проверьте, что файлы загружены в папку examples/")


@dp.errors()
async def error_handler(event: ErrorEvent):
    logger.error(f"Произошла ошибка: {event.exception}")
    
    if hasattr(event.update, 'message') and event.update.message:
        await event.update.message.answer("⚠️ Произошла ошибка. Бот перезапущен.")
        state = dp.fsm.get_context(
            event.update.message.chat.id, 
            event.update.message.from_user.id
        )
        await state.clear()
        await event.update.message.answer("Для какой стадии нужно сделать титульные листы?", reply_markup=stage_keyboard)

async def set_commands(bot: Bot):
    commands = [
        BotCommand(command="start", description="Запустить генерацию титульных листов"),
        BotCommand(command="help", description="Инструкция и примеры файлов"),
    ]
    await bot.set_my_commands(commands)








# --- Обработчик webhook от Telegram ---
async def handle_webhook(request: web.Request):
    try:
        data = await request.json()
        update = Update.model_validate(data)
        await dp.feed_update(bot, update)
    except Exception as e:
        logger.error(f"Ошибка обработки webhook: {e}")
    return web.Response(status=200)

# --- HTTP сервер ---
async def start_webhook_app():
    app = web.Application()
    app.router.add_post(WEBHOOK_PATH, handle_webhook)
    app.router.add_get("/", lambda request: web.Response(text="Bot is running"))
    app.router.add_get("/health", lambda request: web.Response(text="OK"))

    runner = web.AppRunner(app)
    await runner.setup()
    site = web.TCPSite(runner, "0.0.0.0", PORT)
    await site.start()
    logger.info(f"Webhook server running on port {PORT}, path {WEBHOOK_PATH}")
    return runner

# --- Установка webhook у Telegram ---
async def setup_webhook():
    await bot.delete_webhook()
    await bot.set_webhook(url=f"{WEBHOOK_URL}{WEBHOOK_PATH}")
    logger.info(f"Webhook установлен: {WEBHOOK_URL}{WEBHOOK_PATH}")

# --- Основная функция ---
async def main():
    await setup_webhook()
    runner = await start_webhook_app()
    try:
        while True:
            await asyncio.sleep(3600)  # держим сервер работающим
    finally:
        await runner.cleanup()
        await bot.session.close()












# --- Функции обработки таблиц ---
def split_dataframe_PD(df: pd.DataFrame) -> List[pd.DataFrame]:
    subtables = []
    third_col = df.columns[2]
    section_indices = df.index[df[third_col].astype(str).str.contains("Раздел", case=True)].tolist()
    section_indices.append(len(df))

    for i in range(len(section_indices) - 1):
        start_idx, end_idx = section_indices[i], section_indices[i+1]
        section_name = str(df.iloc[start_idx, 2])
        section_df = df.iloc[start_idx+1:end_idx].reset_index(drop=True)

        subsection_indices = section_df.index[
            section_df[third_col].astype(str).str.contains("Подраздел", case=False, na=False)
        ].tolist()
        subsection_indices.append(len(section_df))

        if len(subsection_indices) == 1:
            section_df.columns = ["Том", "Шифр", "Часть"]
            section_df["Раздел"] = section_name
            section_df["Подраздел"] = None
            subtables.append(section_df)
        else:
            for j in range(len(subsection_indices) - 1):
                sub_start_idx, sub_end_idx = subsection_indices[j], subsection_indices[j+1]
                subsection_name = str(section_df.iloc[sub_start_idx, 2])
                subsection_df = section_df.iloc[sub_start_idx+1:sub_end_idx].reset_index(drop=True)

                subsection_df.columns = ["Том", "Шифр", "Часть"]
                subsection_df["Раздел"] = section_name
                subsection_df["Подраздел"] = subsection_name
                subtables.append(subsection_df)

    return subtables


def split_dataframe_RD(df: pd.DataFrame) -> pd.DataFrame:
    second_col = df.columns[1]
    start_idx = df.index[
        df[second_col].astype(str).str.contains("Генеральный план", case=True)].tolist()[0]

    section_df = df.iloc[start_idx:]
    section_df = section_df[section_df[section_df.columns[0]].astype(str).str.len() > 7].reset_index(drop=True)
    section_df.columns = ["Шифр", "Раздел"]

    return section_df


def replace_text_preserve_format_PD(doc: Document, replacements: Dict[str, Optional[str]]) -> None:
    none_keys = {key for key, value in replacements.items() if value is None}
    valid_replacements = {k: str(v) for k, v in replacements.items() if v is not None}

    part_text = str(replacements.get("Название части", ""))
    word_count = len(part_text.split())
    allow_blank_insertion = word_count <= 12

    deleted_count = 0
    cipher_paragraphs = []

    for p in doc.paragraphs:
        for run in p.runs:
            if "Том" in run.text:
                cipher_paragraphs.append(p)
                break

    for p in reversed(doc.paragraphs):
        if any(key in p.text for key in none_keys):
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
            deleted_count += 1
        else:
            for run in p.runs:
                for key, value in valid_replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    if allow_blank_insertion and deleted_count > 0 and cipher_paragraphs:
        for cipher_paragraph in cipher_paragraphs:
            if any(p._element == cipher_paragraph._element for p in doc.paragraphs):
                insert_blank_paragraphs_after(cipher_paragraph, 4)


def replace_text_preserve_format_RD(doc: Document, replacements: Dict[str, Optional[str]]) -> None:

    part_text = str(replacements.get("Название раздела", ""))
    word_count = len(part_text.split())
    allow_blank_insertion = word_count <= 12

    for p in reversed(doc.paragraphs):
        for run in p.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

        if key == "Название раздела" and allow_blank_insertion:
            insert_blank_paragraphs_after(p, 2)




def insert_blank_paragraphs_after(paragraph: Paragraph, count: int) -> None:
    for _ in range(count):
        new_p = OxmlElement('w:p')
        p_pr = OxmlElement('w:pPr')
        new_p.append(p_pr)
        paragraph._element.addnext(new_p)


def create_word_for_each_row_PD(subtables: List[pd.DataFrame], template_path: str, archive_name: str = "documents.zip") -> str:
    temp_files = []

    for _, subtable in enumerate(subtables, start=1):
        for _, row in subtable.iterrows():
            replacements = {
                "Номер": row.get("Том", ""),
                "Название шифра": row.get("Шифр", ""),
                "Название части": row.get("Часть", ""),
                "Название раздела": row.get("Раздел", ""),
                "Название подраздела": row.get("Подраздел", ""),
            }

            doc = Document(template_path)
            replace_text_preserve_format_PD(doc, replacements)

            safe_name = str(row.get("Шифр", "без_шифра")) or "без_шифра"
            safe_name = re.sub(r"[^А-Яа-яA-Za-z0-9_]+", "_", safe_name)

            output_docx = f"{safe_name}.docx"
            doc.save(output_docx)
            temp_files.append(output_docx)

    with zipfile.ZipFile(archive_name, "w") as zipf:
        for f in temp_files:
            zipf.write(f)
            os.remove(f)

    return archive_name


def create_word_for_each_row_RD(subtable: pd.DataFrame, template_path: str, archive_name: str = "documents.zip") -> str:
    temp_files = []

    for _, row in subtable.iterrows():
        replacements = {
            "Название шифра": row.get("Шифр", ""),
            "Название раздела": row.get("Раздел", ""),
        }

        doc = Document(template_path)
        replace_text_preserve_format_RD(doc, replacements)

        safe_name = str(row.get("Шифр", "без_шифра")) or "без_шифра"
        safe_name = re.sub(r"[^А-Яа-яA-Za-z0-9_]+", "_", safe_name)

        output_docx = f"{safe_name}.docx"
        doc.save(output_docx)
        temp_files.append(output_docx)

    with zipfile.ZipFile(archive_name, "w") as zipf:
        for f in temp_files:
            zipf.write(f)
            os.remove(f)

    return archive_name


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("Бот остановлен пользователем")